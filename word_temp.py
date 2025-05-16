import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK
import matplotlib.pyplot as plt
import math
import statistics
import access as acc 
from docx2pdf import convert
import numpy as np
from scipy.stats import binomtest
import sys

def Table_h(table,rows,cols,info,B):
    #********************************************************************************************
    #Description: This creates a table in a word document. 
    #             Input: table object, number of rows, number of columns, info Vector. 
    #             Output: Department instruction sheet. 
    #********************************************************************************************
    # Add a table to the document
    table.alignment = 1  # Center alignment
    # Set table style
    table.style = 'Table Grid'
    # Set the width of the first column
    table.columns[0].width = Inches(1.5)
    # Info vector written to the table.
    k = 0
    for i in range(rows):
        for j in range(cols):
            if k == len(info):
                break
            Cell = table.cell(i,j)
            Cell.text = str(info[k])
            if B == True:
                Cell.paragraphs[0].runs[0].bold = True
            Cell.paragraphs[0].runs[0].font.name = 'Arial' 
            k+=1
    return table

def Table_v(table, rows, cols, info, B):
    #********************************************************************************************
    #Description: This creates a table in a word document. 
    #             Input: table object, number of rows, number of columns, info Vector. 
    #             Output: Department instruction sheet. 
    #********************************************************************************************
    # Add a table to the document
    table.alignment = 1  # Center alignment
    # Set table style
    table.style = 'Table Grid'
    # Set the width of the first column
    table.columns[0].width = Inches(1.5)
    # Info vector written to the table.
    k = 0
    for j in range(cols):  # Loop through columns first
        for i in range(rows):  # Then loop through rows
            if k == len(info):
                break
            Cell = table.cell(i, j)  # Note the reversed order of i and j
            Cell.text = str(info[k])
            if B == True:
                Cell.paragraphs[0].runs[0].bold = True
            Cell.paragraphs[0].runs[0].font.name = 'Arial'
            k += 1
    return table

def sample_std_dev(data):
    n = len(data)  # number of samples less than 30
    mean = sum(data) / n
    deviations = [(x - mean) ** 2 for x in data]
    variance = sum(deviations) / (n - 1)
    std_dev = math.sqrt(variance)
    return std_dev

def sample_variance(data):
    n = len(data)
    mean = sum(data) / n
    deviations = [(x - mean) ** 2 for x in data]
    variance = sum(deviations) / (n - 1)
    return variance

def flatten(lst):
    flattened_list = []
    for i in lst:
        if isinstance(i, list):
            flattened_list.extend(flatten(i))
        else:
            flattened_list.append(i)
    return flattened_list

def floater(vec):
    new_vec = []
    for i in range(len(vec)):
        new_vec.append(float(vec[i]))
    return new_vec

def analysis(vec):
    ran = max(vec)-min(vec)
    avg = statistics.mean(vec)
    std = sample_std_dev(vec)
    variance = sample_variance(vec)
    vec = [avg,ran,std,variance]  
    return vec

def sig_fig(data):
    mod_data = []
    for i in range(len(data)):
        mod_data.append(format(data[i],'.3f'))
    return mod_data

def sig_fig_2(data):
    mod_data = []
    for i in range(len(data)):
        mod_data.append(format(data[i],'.4f'))
    return mod_data

def binomial(k):
    n = 10  # Number of samples measured 
    # Hypothesized proportion of defectives (acceptable defect rate)
    p0 = 0.001  # Example: 1% is the acceptable defect rate

    # Perform the binomial test using binomtest
    result = binomtest(k, n, p0, alternative='greater')

    # Display the result
    # print(f"P-value: {result.pvalue:.4f}")

    # Determine if the result is statistically significant
    alpha = 0.05  # Significance level
    if result.pvalue < alpha:
        return True # Sorting required
    else:
        return False # Sorting Not Required 

def auto_disposition(part_info_vec,data, spec, key,diel,passivation):
    min,max = 0,0
    if (key == 'thickness') and (part_info_vec[0][-1] != 'G'):
        spec[1] = spec[1] + (2/1000)
    
    if (key == 'thickness') and (part_info_vec[0][-1] == 'G') and (passivation == True):
        spec[1] = spec[1] - (5/1000)

    for i in range(len(data)):
        if float(data[i])<float(spec[0]):
            min+=1
        elif float(data[i])>float(spec[1]):
            max+=1
            
    under = binomial(min)
    over = binomial(max)
    
    text = ''
    if under == True and over == True:
        text = f'Send parts to MRB! Multiple dimentional defects with {key}.'
    else:
        if key == 'OD/L':
            if under == True:
                if part_info_vec[0][-1] != 'G' and diel != 'NPO':
                    yield_loss = (min/10)*100
                    text = text + f'See engineering at termination for 2X OD paint. Also, contact PC for potential {yield_loss}% yield loss. '
                else:
                    text = text + f'Send parts to MRB. Issue {key}. Spec: {spec[0]}" to {spec[1]}"'
            elif over == True:
                text = text + f'2X Tumble. Issue {key}. Spec: {spec[0]}" to {spec[1]}"'
            else:
                text = f'Continue. No Problem Found with {key}. Spec: {spec[0]}" to {spec[1]}"'
        elif key == 'thickness':
            if under == True:
                if passivation == True:
                    text = text + f'2X Passivation for {key}. Spec: {spec[0]}" to {spec[1]}"'
                else:
                    text = text + f'Send to MRB. Issue {key}, Spec: {spec[0]}" to {spec[1]}"'
            elif over == True:
                if part_info_vec[0][-1] != 'G':
                    grnd_spec = spec[1] - (5/1000)
                    grnd_min = str(format(grnd_spec,'.4f'))
                    grnd_max = str(format(grnd_spec+(0.5/1000),'.4f'))
                    text = text + f'Send parts for polishing. Specification: {grnd_min}" to {grnd_max}"'
                elif (part_info_vec[0][-1] == 'G') and (passivation == True):
                        grnd_spec = spec[1] - (5/1000)
                        grnd_min = str(format(grnd_spec,'.4f'))
                        grnd_max = str(format(grnd_spec+(0.5/1000),'.4f'))
                        text = text + f'Send parts to Polishing. Specification: {grnd_min}" to {grnd_max}"'    
            else:
                text = f'Continue. No Problem Found with {key}. Spec: {spec[0]}" to {spec[1]}"'
        elif key =='Warpage':
            if over == True:
                text = f'Flat-Fire the parts. Issue {key}. Spec: {spec[0]}" to {spec[1]}"'
            else:
                text = f'Continue. No Problem Found with {key}. Spec: {spec[0]}" to {spec[1]}"'
    
    return text



    # # Print results
    # print(f"Proportion of parts out of spec: {proportion_out_of_spec:.2f}")
    # print(f"Expected proportion under null hypothesis: {expected_proportion:.2f}")
    # print(f"P-value for hypothesis test: {p_value:.4f}")

    # # Make decision based on p-value
    # if p_value < alpha:
    #     if key == 'OD/L':
    #         text = f'If {key} .'
    # else:
    #     text = f'{key} in spec. '
    #     return text
def Plot(vec1, vec2, title1, title2, ylabel,path):
    def plot_control_chart(ax, data, title, ylabel):
        mean = statistics.mean(data)
        sigma = sample_std_dev(data)  # Calculate sample standard deviation
        upper_control_limit = mean + 3 * sigma
        lower_control_limit = mean - 3 * sigma

        # Calculate ±1 sigma and ±2 sigma limits
        one_sigma_upper = mean + sigma
        one_sigma_lower = mean - sigma
        two_sigma_upper = mean + 2 * sigma
        two_sigma_lower = mean - 2 * sigma

        # Plot the control chart
        ax.plot(data, marker='o', linestyle='-', color='b', label='Data')
        ax.axhline(upper_control_limit, color='r', linestyle='--', label='Upper Control Limit (3\u03C3)')
        ax.axhline(lower_control_limit, color='r', linestyle='--', label='Lower Control Limit (3\u03C3)')
        ax.axhline(mean, color='g', linestyle='-', label='Mean')

        # ±1 sigma lines
        ax.axhline(one_sigma_upper, color='orange', linestyle='dashdot', label='Upper Limit (1\u03C3)')
        ax.axhline(one_sigma_lower, color='orange', linestyle='dashdot', label='Lower Limit (1\u03C3)')

        # ±2 sigma lines
        ax.axhline(two_sigma_upper, color='purple', linestyle='dotted', label='Upper Limit (2\u03C3)')
        ax.axhline(two_sigma_lower, color='purple', linestyle='dotted', label='Lower Limit (2\u03C3)')

        ax.set_title(title)
        ax.set_xlabel('Samples')
        ax.set_ylabel(ylabel)
        ax.legend()
        ax.grid(True)
    
    # Create figure and subplots
    fig, (ax1, ax2) = plt.subplots(nrows=1, ncols=2, figsize=(15, 6))

    # Plot control chart for vec1
    plot_control_chart(ax1, vec1, title1, ylabel)

    # Plot control chart for vec2
    plot_control_chart(ax2, vec2, title2, ylabel)

    plt.tight_layout()
    # Save the chart as an image
    plt.savefig(path)
    plt.close()  # Close the plot to prevent displaying it

def word_drill(part_info_vec, frame_ID): 

    #********************************************************************************************
    #Description: This function checks the connection to the Access Database. 
    #             Input: vector containing info from deparment functions.
    #             Output: Department instruction sheet for Lamination and Drilling. 
    #********************************************************************************************
    directory = r"K:\Public\DATA SWIFT\Word Files\Drill"
    #---------------------Data Manip----------------------------------
    data = acc.read_drill_database(part_info_vec)
    
    #------------------------------------------------------------------
    data = flatten(data)
    OD_L = floater(data[3:13])
    Width = floater(data[13:23])
    Thickness_w = floater(data[23:33])
    Thickness_wo = floater(data[33:43])
    ID1 = floater(data[43:53])
    ID2 = floater(data[53:63])
    ID3 = floater(data[63:73])
    Warpage = floater(data[73:83])
    #-------------------------------------------------------------------
    # Analysis
    OD_L_a = analysis(OD_L)
    Width_a = analysis(Width)
    Thickness_w_a = analysis(Thickness_w)
    Thickness_wo_a = analysis(Thickness_wo)
    ID1_a = analysis(ID1)
    ID2_a = analysis(ID2)
    ID3_a = analysis(ID3)
    Warpage_a = analysis(Warpage)
    Analysis = []
    #-------------------------------------------------------------------

    frame_ID.pop()
    new_data = []
    specs = []
    row_names = ['Part']
    row_names_a = ['Type']
    for i in range(len(frame_ID)):
        if frame_ID[i] == 'OD' or frame_ID[i] == 'Length':
            new_data.append(OD_L)
            Analysis.append(OD_L_a)
            if frame_ID[i] == 'OD':
                row_names.append('OD (In)')
                row_names_a.append('OD')
            elif frame_ID[i] == 'Length':
                row_names.append('Length (In)')
                row_names_a.append('Length (In)')
        elif frame_ID[i] == 'Width':
            new_data.append(Width)
            Analysis.append(Width_a)
            row_names.append('Width (In)')
            row_names_a.append('Width')
        elif frame_ID[i] == 'Thickness\nWith Top Layer':
            new_data.append(Thickness_w)
            Analysis.append(Thickness_w_a)
            row_names.append('Height 1 (In)')
            row_names_a.append('Height 1')
        elif frame_ID[i] == 'Thickness\nWithout Top Layer':
            new_data.append(Thickness_wo)
            Analysis.append(Thickness_wo_a)
            row_names.append('Height 2 (In)')
            row_names_a.append('Height 2')
        elif frame_ID[i] == 'ID A':
            new_data.append(ID1)
            Analysis.append(ID1_a)
            row_names.append('ID A (In)')
            row_names_a.append('ID A')
        elif frame_ID[i] == 'ID B':
            new_data.append(ID2)
            Analysis.append(ID2_a)
            row_names.append('ID B (In)')
            row_names_a.append('ID B')
        elif frame_ID[i] == 'ID C':
            new_data.append(ID3)
            Analysis.append(ID3_a)
            row_names.append('ID C (In)')
            row_names_a.append('ID C')
        elif frame_ID[i] == 'Warpage':
            new_data.append(Warpage)
            Analysis.append(Warpage_a)
            row_names.append('Warpage (In)')
            row_names_a.append('Warpage')
    
    index = [1,2,3,4,5,6,7,8,9,10]
    table_data = sig_fig(flatten(new_data))
    table_data = flatten(index + table_data)
    
    index = ['Avg (In)','Range (In)','\u03C3 (In)','\u03C3\u00B2 (In\u00B2)']
    table_analysis = sig_fig_2(flatten(Analysis))
    table_analysis = flatten(index + table_analysis)
    #===============================================================================================================================================
    # Create a new Word document
    doc = Document()
    # Add a header to the document
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Drill Data Summary"
    # Set header text alignment to center
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set header font size to 28 and make it bold
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(28)
    header_run.bold = True
    header_run.bold = True

    #****************Part Info*********************************
    # Set header font size to 28 and make it bold
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(28)
    header_run.bold = True
    header_run.bold = True
    part_info = ['Part Number','Revision','MO Number','Document Number',part_info_vec[0],part_info_vec[2],part_info_vec[1],'4215P023 Rev.0']
    # Doccument info: 
    # Add a table to the document
    rows=2
    cols=4
    B = False
    table = doc.add_table(rows, cols)
    table = Table_h(table,rows,cols,part_info,B)

    #**********************************************************
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Measurements")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True

    #================Measurements Table=======================
    rows = 1
    cols = len(frame_ID) + 1
    B = True
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,row_names,B)

    rows = 10
    cols = len(frame_ID) + 1
    B = False
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,table_data,B)
    

    #**********************************************************
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Statistical Analysis")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True

    #================Analysis Table=======================
    rows = 1
    cols = len(frame_ID) + 1
    B = True
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,row_names_a,B)

    rows = 4
    cols = len(frame_ID) + 1
    B = False
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,table_analysis,B)

    #================Add Plots=============================
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Plots")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True 

    path =  r"K:\Public\DATA SWIFT\IMG\Drill.png"
    Plot(OD_L,Thickness_wo,'OD\L Measurements','Height Measurements','Inches',path)

    # Add the bar chart image to the Word document
    # Add a paragraph with center alignment
    center_aligned_paragraph = doc.add_paragraph()
    center_aligned_paragraph.alignment = 1  # 1 for center alignment

    # Add the picture to the center-aligned paragraph
    center_aligned_run = center_aligned_paragraph.add_run()
    center_aligned_run.add_picture(path, height = Inches(3.3), width=Inches(6))
    # doc.add_picture(chart_image_path, width=Inches(5))

    # auto Dispostion
    # text = auto_disposition_drill(OD_L,Thickness_wo,specs)
    # Specify the directory and filename
    filename = f'Drill Report {part_info_vec[0]}_{part_info_vec[1]}_{part_info_vec[2]}.docx'
    save_path = os.path.join(directory, filename)

    # Create the directory if it doesn't exist
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Save the document at the specified location
    doc.save(save_path)

    return save_path

def word_FT(part_info_vec, frame_ID,FT_specs): 

    #********************************************************************************************
    #Description: This function checks the connection to the Access Database. 
    #             Input: vector containing info from deparment functions.
    #             Output: Department instruction sheet for Lamination and Drilling. 
    #********************************************************************************************
    directory = r"K:\Public\DATA SWIFT\Word Files\FT"
    #---------------------Data Manip----------------------------------
    data = acc.read_FT_database(part_info_vec)
    ceramic = acc.ceramic_info(part_info_vec[0],part_info_vec[2]) # Returns 'TCC','CERAMIC_TY','CERAMIC_LO'
    #------------------------------------------------------------------
    data = flatten(data)
    OD_L = floater(data[3:13])
    Width = floater(data[13:23])
    Thickness = floater(data[23:33])
    ID1 = floater(data[33:43])
    ID2 = floater(data[43:53])
    ID3 = floater(data[53:63])
    Warpage = floater(data[63:73])
    #-------------------------------------------------------------------
    # Analysis
    OD_L_a = analysis(OD_L)
    Width_a = analysis(Width)
    Thickness_a = analysis(Thickness)
    ID1_a = analysis(ID1)
    ID2_a = analysis(ID2)
    ID3_a = analysis(ID3)
    Warpage_a = analysis(Warpage)
    Analysis = []
    #-------------------------------------------------------------------
    #----------------From Drill DB--------------------------------------
    #---------------------Data Manip----------------------------------
    data_d = acc.read_drill_database(part_info_vec)
    data_d = flatten(data_d)
    #------------------------------------------------------------------
    OD_L_d = floater(data_d[3:13])
    Width_d = floater(data_d[13:23])
    Thickness_wo_d = floater(data_d[33:43])
    ID1_d = floater(data_d[43:53])
    ID2_d = floater(data_d[53:63])
    ID3_d = floater(data_d[63:73])
    #-------------------------------------------------------------------
    #-------------------------------------------------------------------
    frame_ID.pop()
    new_data = []
    specs = []
    row_names = ['Part']
    row_names_a = ['Type']
    for i in range(len(frame_ID)):
        if frame_ID[i] == 'OD' or frame_ID[i] == 'Length':
            new_data.append(OD_L)
            Analysis.append(OD_L_a)
            if frame_ID[i] == 'OD':
                row_names.append('OD (In)')
                row_names_a.append('OD')
                specs.append([float(FT_specs[2]),float(FT_specs[3])])
            elif frame_ID[i] == 'Length':
                row_names.append('Length (In)')
                row_names_a.append('Length (In)')
                specs.append([float(FT_specs[2]),float(FT_specs[3])])
        elif frame_ID[i] == 'Width':
            new_data.append(Width)
            Analysis.append(Width_a)
            row_names.append('Width (In)')
            row_names_a.append('Width')
        elif frame_ID[i] == 'Thickness':
            new_data.append(Thickness)
            Analysis.append(Thickness_a)
            row_names.append('Height (In)')
            row_names_a.append('Height')
            specs.append([float(FT_specs[6]),float(FT_specs[7])])
        elif frame_ID[i] == 'ID A':
            new_data.append(ID1)
            Analysis.append(ID1_a)
            row_names.append('ID A (In)')
            row_names_a.append('ID A')
        elif frame_ID[i] == 'ID B':
            new_data.append(ID2)
            Analysis.append(ID2_a)
            row_names.append('ID B (In)')
            row_names_a.append('ID B')
        elif frame_ID[i] == 'ID C':
            new_data.append(ID3)
            Analysis.append(ID3_a)
            row_names.append('ID C (In)')
            row_names_a.append('ID C')
        elif frame_ID[i] == 'Warpage':
            new_data.append(Warpage)
            Analysis.append(Warpage_a)
            row_names.append('Warpage (In)')
            row_names_a.append('Warpage')
            specs.append([0,float(FT_specs[14])])
    
    specs = flatten(specs)
    index = [1,2,3,4,5,6,7,8,9,10]
    table_data = sig_fig(flatten(new_data))
    table_data = flatten(index + table_data)
    
    index = ['Avg (In)','Range (In)','\u03C3 (In)','\u03C3\u00B2 (In\u00B2)']
    table_analysis = sig_fig_2(flatten(Analysis))
    table_analysis = flatten(index + table_analysis)
    #===============================================================================================================================================
    # Create a new Word document
    doc = Document()
    # Add a header to the document
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Fire-Test Data Summary"
    # Set header text alignment to center
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set header font size to 28 and make it bold
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(28)
    header_run.bold = True
    header_run.bold = True

    #****************Part Info*********************************
    part_info = ['Part Number','Revision','MO Number','Document Number',part_info_vec[0],part_info_vec[2],part_info_vec[1],'4216P012 Rev. 0']
    # Doccument info: 
    # Add a table to the document
    rows=2
    cols=4
    B = False
    table = doc.add_table(rows, cols)
    table = Table_h(table,rows,cols,part_info,B)

    #****************Ceramic Info*********************************
    # Set header font size to 28 and make it bold
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Ceramic Information")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True
    if ceramic[1] == 'UL990':
        ceramic[1] = 'ULF990'
    part_info = ['Ceramic Type','Ceramic Name','Ceramic Lot',ceramic[0],ceramic[1],ceramic[2]]
    # Doccument info: 
    # Add a table to the document
    rows=2
    cols=3
    B = False
    table = doc.add_table(rows, cols)
    table = Table_h(table,rows,cols,part_info,B)

    #**********************************************************
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Measurements")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True

    #================Measurements Table=======================
    rows = 1
    cols = len(frame_ID) + 1
    B = True
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,row_names,B)

    rows = 10
    cols = len(frame_ID) + 1
    B = False
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,table_data,B)
    

    #**********************************************************
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Statistical Analysis")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True

    #================Analysis Table=======================
    rows = 1
    cols = len(frame_ID) + 1
    B = True
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,row_names_a,B)

    rows = 4
    cols = len(frame_ID) + 1
    B = False
    table = doc.add_table(rows, cols)
    table = Table_v(table,rows,cols,table_analysis,B)

    #================Add Plots=============================
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Measurement Plots")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True 

    # Calculate standard deviation for each dataset
    path =  r"K:\Public\DATA SWIFT\IMG\FT_1.png"
    Plot(OD_L,Thickness,'OD\L Measurements','Height Measurements','Inches',path)

    # Add the bar chart image to the Word document
    # Add a paragraph with center alignment
    center_aligned_paragraph = doc.add_paragraph()
    center_aligned_paragraph.alignment = 1  # 1 for center alignment

    # Add the picture to the center-aligned paragraph
    center_aligned_run = center_aligned_paragraph.add_run()
    center_aligned_run.add_picture(path, height = Inches(3.1), width=Inches(6))
    # doc.add_picture(chart_image_path, width=Inches(5))
    #----------------Page 2------------------------------------------------------
    # Add a page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Shrinkage Analysis")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True

    # Shrinkage Table----------------------------
    avg_OD_L_d = statistics.mean(OD_L_d)
    avg_Thickness_wo_d = statistics.mean(Thickness_wo_d)
    od_shrinkage = 100*((avg_OD_L_d - statistics.mean(OD_L))/(avg_OD_L_d))
    height_shrinkage = 100*((avg_Thickness_wo_d - statistics.mean(Thickness))/(avg_Thickness_wo_d))
        # Calculate mean of standard deviations
    
    mean_std_dev1 =  sample_std_dev(OD_L)
    mean_std_dev2 = sample_std_dev(Thickness)

    # Format the values for display
    formatted_od_shrinkage = f'{od_shrinkage:.3f}'
    formatted_height_shrinkage = f'{height_shrinkage:.3f}'
    formatted_mean_std_dev1 = f'{mean_std_dev1:.3f}'
    formatted_mean_std_dev2 = f'{mean_std_dev2:.3f}'

    # Construct the part_info list with formatted values
    part_info = [
        'OD Shrinkage (%)',
        'Height Shrinkage (%)',
        f'{formatted_od_shrinkage} \u00B1 {formatted_mean_std_dev1}',
        f'{formatted_height_shrinkage} \u00B1 {formatted_mean_std_dev2}']

    # Doccument info: 
    # Add a table to the document
    rows=2
    cols=2
    B = False
    table = doc.add_table(rows, cols)
    table = Table_h(table,rows,cols,part_info,B)

      #================Add Plots=============================
    OD_shrink = []
    Thickness_shrink = []
    
    for i in range(len(OD_L)):
        OD_shrink.append(100*((avg_OD_L_d - OD_L[i]))/(avg_OD_L_d))
    
    for j in range(len(Thickness)):
        Thickness_shrink.append(100*((avg_Thickness_wo_d - Thickness[j])/(avg_Thickness_wo_d)))
    # Calculate standard deviation for each dataset
   
    # Add a subheading for Measured Data at Drill
    subheading = doc.add_paragraph("Shrinkage Plots")
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subheading_run = subheading.runs[0]
    subheading_run.font.size = Pt(16)
    subheading_run.bold = True 

     # Calculate standard deviation for each dataset
    path =  r"K:\Public\DATA SWIFT\IMG\FT_2.png"
    Plot(OD_shrink,Thickness_shrink,'OD\L Shrinkage','Height Shrinkage','Percent (%)',path)

    # Add the bar chart image to the Word document
    # Add a paragraph with center alignment
    center_aligned_paragraph = doc.add_paragraph()
    center_aligned_paragraph.alignment = 1  # 1 for center alignment

    # Add the picture to the center-aligned paragraph
    center_aligned_run = center_aligned_paragraph.add_run()
    center_aligned_run.add_picture(path, height = Inches(4), width=Inches(6))

    # auto Dispostion
    key = 'OD/L'
    spec = [specs[0],specs[1]]
    diel = ceramic[0]
    passivation = acc.passivation_check(part_info_vec)
    text_1 = auto_disposition(part_info_vec,OD_L, spec, key,diel,passivation)
    key = 'thickness'
    spec = [specs[2],specs[3]]
    text_2 = auto_disposition(part_info_vec,Thickness, spec, key,diel,passivation)
    
    if frame_ID[-1] == 'Warpage':
        spec = [specs[4],specs[5]]
        
        key = 'Warpage'
        text_3 = auto_disposition(part_info_vec,Warpage, spec, key,diel,passivation)
        paragraph = doc.add_paragraph(f"Auto Disposition:\n\t{text_1}\n\t{text_2}\n\t{text_3}")
        paragraph = paragraph.runs[0]
        paragraph.bold = True
        paragraph.font.size = Pt(12)
    else:
        paragraph = doc.add_paragraph(f"Auto Disposition:\n\t{text_1}\n\t{text_2}")
        paragraph = paragraph.runs[0]
        paragraph.bold = True
        paragraph.font.size = Pt(12)


    # Specify the directory and filename
    filename = f'Fire-test Report {part_info_vec[0]}_{part_info_vec[1]}_{part_info_vec[2]}.docx'
    save_path = os.path.join(directory, filename)

    # Create the directory if it doesn't exist
    if not os.path.exists(directory):
        os.makedirs(directory)

    # Save the document at the specified location
    doc.save(save_path)

    return save_path

def create_pdf(filename,location):
    if location == 'Drill':
        output_dir = r'K:\Public\DATA SWIFT\Drill Reports'
    elif location == 'FT':
        output_dir = r'K:\Public\DATA SWIFT\Fire Test Reports'
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Specify the output file path for the PDF
    sys.stderr = open("consoleoutput.log", "w") # to prevent consol progress bar
    output_filename = os.path.join(output_dir, os.path.splitext(os.path.basename(filename))[0] + ".pdf")

    # Convert the DOCX file to PDF and save it to the specified output file path
    convert(filename, output_filename)
    
# part_info_vec = ['55CA215-B','430857-00','NCC-2-ENG']
# frame_ID = ['OD','Thickness','ID A','Warpage','Review']
# FT_specs = ['pn','rev','1.163','1.173','0.0','0.0','0.09','0.10','0.051','0.057','0.0','0.0','0.0','0.0','0.005']
# word_FT(part_info_vec, frame_ID,FT_specs)